"""Format conversion helpers."""

import os
import fitz
from PIL import Image


class Converter:
    """PDF <-> images / text / docx (optional)."""

    # ---- PDF pages to images ----
    @staticmethod
    def pdf_to_images(pdf_path: str, out_folder: str, fmt: str = "png",
                      dpi: int = 150, pages: list[int] | None = None,
                      progress_cb=None) -> list[str]:
        os.makedirs(out_folder, exist_ok=True)
        doc = fitz.open(pdf_path)
        zoom = dpi / 72.0
        mat = fitz.Matrix(zoom, zoom)
        page_indices = pages if pages else list(range(doc.page_count))
        total = len(page_indices)
        outputs = []
        ext = "jpg" if fmt.lower() in ("jpg", "jpeg") else "png"
        for i, idx in enumerate(page_indices, 1):
            page = doc.load_page(idx)
            pix = page.get_pixmap(matrix=mat, alpha=False)
            out = os.path.join(out_folder, f"page_{idx + 1:04d}.{ext}")
            pix.save(out)
            outputs.append(out)
            if progress_cb:
                progress_cb(i, total)
        doc.close()
        return outputs

    # ---- images to PDF ----
    @staticmethod
    def images_to_pdf(image_paths: list[str], output_path: str):
        if not image_paths:
            raise ValueError("No images provided")

        def _open_pdf_safe_image(path: str) -> Image.Image:
            """Open an image and convert it safely for PDF output.

            Important: simply calling convert("RGB") on a transparent PNG drops
            the alpha channel onto a black background. This made transparent PNGs
            look black after Image→PDF conversion. We flatten images with alpha
            onto a white page instead, which is the expected PDF-paper look.
            """
            img = Image.open(path)
            if img.mode in ("RGBA", "LA") or (img.mode == "P" and "transparency" in img.info):
                rgba = img.convert("RGBA")
                white = Image.new("RGBA", rgba.size, (255, 255, 255, 255))
                white.alpha_composite(rgba)
                img.close()
                return white.convert("RGB")
            return img.convert("RGB")

        images = []
        try:
            for p in image_paths:
                images.append(_open_pdf_safe_image(p))
            first, rest = images[0], images[1:]
            first.save(output_path, "PDF", resolution=100.0,
                       save_all=True, append_images=rest)
        finally:
            for img in images:
                img.close()

    # ---- extract text ----
    @staticmethod
    def pdf_to_text(pdf_path: str, output_path: str | None = None,
                    pages: list[int] | None = None) -> str:
        doc = fitz.open(pdf_path)
        page_indices = pages if pages else list(range(doc.page_count))
        chunks = []
        for idx in page_indices:
            chunks.append(doc.load_page(idx).get_text())
        doc.close()
        full = "\n\n".join(chunks)
        if output_path:
            with open(output_path, "w", encoding="utf-8") as f:
                f.write(full)
        return full

    # ---- to DOCX / Word ----
    @staticmethod
    def pdf_to_docx(pdf_path: str, output_path: str, progress_cb=None) -> bool:
        """Best-effort PDF to Word conversion.

        This creates a real .docx file using python-docx. It preserves readable
        text page-by-page and inserts page breaks between PDF pages. For scanned
        image-only PDFs, the resulting Word file may be empty unless the user runs
        OCR first, because normal PDF text extraction cannot read pixels.

        Returns True on success and False when python-docx is not installed.
        """
        try:
            from docx import Document
            from docx.shared import Pt
        except ImportError:
            return False

        docx = Document()
        styles = docx.styles
        styles["Normal"].font.name = "Arial"
        styles["Normal"].font.size = Pt(10)

        pdf = fitz.open(pdf_path)
        total = max(pdf.page_count, 1)
        try:
            for page_no in range(pdf.page_count):
                if progress_cb:
                    progress_cb(page_no + 1, total)

                page = pdf.load_page(page_no)
                if pdf.page_count > 1:
                    docx.add_paragraph(f"Page {page_no + 1}").style = styles["Heading 2"]

                text = page.get_text("text") or ""
                paragraphs = [p.strip() for p in text.split("\n") if p.strip()]
                if paragraphs:
                    for paragraph in paragraphs:
                        docx.add_paragraph(paragraph)
                else:
                    docx.add_paragraph(
                        "[No selectable text found on this page. Run OCR first for scanned PDFs.]"
                    )

                if page_no < pdf.page_count - 1:
                    docx.add_page_break()
        finally:
            pdf.close()

        docx.save(output_path)
        return True
